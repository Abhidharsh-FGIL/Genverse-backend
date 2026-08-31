"""
ebook_export_service.py

Server-side PDF (ReportLab Platypus) and DOCX (python-docx) generation
for professional book-format export of eBooks.

Layout per user spec:
  Page 1  — Cover (title, author, image, description)
  Page 2  — About This Book
  Page 3  — Table of Contents  (with dotted leaders + estimated page nums)
  Page 4+ — Chapters  (each starts on a fresh page)
             eyebrow | 22pt title | rule | image | 12pt justified body |
             image | key-points box | page number
  Final   — Assessment Questions  (if present)
           — Thank You  (if present)

Typography:
  Font      : Times New Roman / Times-Roman  (12pt body, 22pt chapter titles)
  Line space: 1.5 × font size
  Margins   : 1 inch on all sides  (handled by ReportLab + python-docx)
  Images    : centred, aspect-ratio preserved, never stretched
"""

from __future__ import annotations

import base64
import io
import os
import re
from xml.sax.saxutils import escape as xml_escape

from PIL import Image as PILImage

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    HRFlowable,
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

# ── python-docx ───────────────────────────────────────────────────────────────
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ── Geometry constants ────────────────────────────────────────────────────────
PAGE_W, PAGE_H = A4               # 595.28 pt × 841.89 pt
MARGIN = 1.0 * inch               # 72 pt — 1-inch margins everywhere
CONTENT_W = PAGE_W - 2 * MARGIN   # ~451 pt  (usable width in PDF)

A4_W_IN   = 8.27                  # A4 width in inches (DOCX)
DOC_MARGIN = 1.0                  # inches

FONT_RL = "Times-Roman"           # ReportLab built-in
FONT_DOCX = "Times New Roman"     # python-docx

# ── Non-Latin script fonts (PDF only) ───────────────────────────────────────────
# ReportLab's built-in "Times-*" fonts are the base-14 PDF standard fonts, which
# only cover WinAnsi/Latin-1 — any Tamil/Devanagari/etc. character silently draws
# as nothing (a "tofu" box) with no embedded Unicode font registered. DOCX/in-app
# viewing were never affected by this because those rely on the OS's or browser's
# own installed fonts, not a font baked into the exported file — a PDF has no such
# fallback, so the correct font must be embedded explicitly at generation time.
_FONTS_DIR = os.path.join(os.path.dirname(__file__), "..", "assets", "fonts")
_SCRIPT_FONT_FILES = {
    "hi": ("NotoSansDevanagari", "NotoSansDevanagari.ttf"),
    "ta": ("NotoSansTamil", "NotoSansTamil.ttf"),
    "te": ("NotoSansTelugu", "NotoSansTelugu.ttf"),
    "kn": ("NotoSansKannada", "NotoSansKannada.ttf"),
    "ml": ("NotoSansMalayalam", "NotoSansMalayalam.ttf"),
}
_registered_script_fonts: set[str] = set()


def _fonts_for_language(language: str | None) -> tuple[str, str, str]:
    """Return (regular, bold, italic) ReportLab font names for the ebook's language.

    English (or any unrecognised code) keeps the existing Times-Roman/-Bold/-Italic
    trio unchanged. For the other codes this platform supports (see LANGUAGE_LOCALE
    in news_service.py), a matching Noto Sans <Script> TTF is registered on first
    use and reused for regular/bold/italic alike — these are variable-weight font
    files but ReportLab doesn't interpret variable-font axes, so it only ever draws
    the single default (Regular) instance regardless of which of the three names is
    requested. A non-bold Tamil heading is a minor cosmetic gap; a blank one is not.
    """
    entry = _SCRIPT_FONT_FILES.get((language or "en").lower())
    if not entry:
        return "Times-Roman", "Times-Bold", "Times-Italic"

    font_name, file_name = entry
    if font_name not in _registered_script_fonts:
        pdfmetrics.registerFont(TTFont(font_name, os.path.join(_FONTS_DIR, file_name)))
        pdfmetrics.registerFontFamily(
            font_name, normal=font_name, bold=font_name, italic=font_name, boldItalic=font_name,
        )
        _registered_script_fonts.add(font_name)
    return font_name, font_name, font_name


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

def _decode_image(data_url: str) -> bytes | None:
    """Return raw bytes from a base64 data URL, or None on failure."""
    if not data_url:
        return None
    try:
        m = re.match(r"data:[^;]+;base64,(.+)", data_url, re.DOTALL)
        return base64.b64decode(m.group(1)) if m else None
    except Exception:
        return None


def _pdf_text(value: object) -> str:
    """Escape dynamic text for ReportLab Paragraph XML parser."""
    return xml_escape("" if value is None else str(value))


# ─────────────────────────────────────────────────────────────────────────────
# Rich-content parsing (shared by PDF + DOCX)
#
# The AI is instructed (see ai_service.py's chapter-generation prompt) to format
# chapter/summary text with a small set of HTML tags: <h3> sub-headings, <b>/<i>
# emphasis, <ul>/<ol>/<li> lists, and <pre><code> blocks. The in-app viewer
# parses this into real styled elements; without the same parsing here, those
# tags used to show up as literal text (e.g. "<h3>Some heading</h3>") in the
# downloaded PDF/DOCX. These helpers split raw content into typed blocks and
# inline (text, bold, italic) runs so both exporters can render it properly —
# any *other* stray tag the model emits (e.g. an over-generated <h1>) is simply
# dropped rather than shown as literal text.
# ─────────────────────────────────────────────────────────────────────────────

_BLOCK_RE = re.compile(
    r"<h[1-6]>.*?</h[1-6]>|<ul>.*?</ul>|<ol>.*?</ol>|<pre>\s*<code[^>]*>.*?</code>\s*</pre>",
    re.IGNORECASE | re.DOTALL,
)
_LI_RE = re.compile(r"<li>(.*?)</li>", re.IGNORECASE | re.DOTALL)
_INLINE_TAG_RE = re.compile(r"<(/?)(b|i)>", re.IGNORECASE)
_ANY_TAG_RE = re.compile(r"<[^>]+>")


def _parse_content_blocks(text: str) -> list[dict]:
    """Split raw AI-generated text into typed blocks: paragraph, heading,
    (un)ordered list, or code — instead of treating <h3>/<ul>/<pre> as plain text."""
    if not text:
        return []
    blocks: list[dict] = []
    pos = 0
    for m in _BLOCK_RE.finditer(text):
        if m.start() > pos:
            blocks += [{"type": "para", "text": p} for p in _split_paras(text[pos:m.start()])]
        raw = m.group(0)
        h_m = re.match(r"<h[1-6]>(.*?)</h[1-6]>", raw, re.IGNORECASE | re.DOTALL)
        if h_m:
            blocks.append({"type": "heading", "text": h_m.group(1).strip()})
        else:
            code_m = re.match(r"<pre>\s*<code[^>]*>(.*?)</code>\s*</pre>", raw, re.IGNORECASE | re.DOTALL)
            if code_m:
                blocks.append({"type": "code", "text": code_m.group(1).strip("\n")})
            else:
                list_m = re.match(r"<(ul|ol)>(.*?)</\1>", raw, re.IGNORECASE | re.DOTALL)
                if list_m:
                    blocks.append({
                        "type": "list",
                        "ordered": list_m.group(1).lower() == "ol",
                        "items": [li.strip() for li in _LI_RE.findall(list_m.group(2)) if li.strip()],
                    })
        pos = m.end()
    if pos < len(text):
        blocks += [{"type": "para", "text": p} for p in _split_paras(text[pos:])]
    return blocks


def _split_inline_runs(text: str) -> list[tuple[str, bool, bool]]:
    """Parse a line of text into (text, bold, italic) runs, honoring <b>/<i> tags
    and silently dropping any other stray HTML tag rather than showing it literally."""
    runs: list[tuple[str, bool, bool]] = []
    bold = italic = False
    pos = 0
    tag_re = re.compile(r"<(/?)(b|i)>|<[^>]+>", re.IGNORECASE)
    for m in tag_re.finditer(text):
        if m.start() > pos:
            runs.append((text[pos:m.start()], bold, italic))
        if m.group(2):
            closing = bool(m.group(1))
            if m.group(2).lower() == "b":
                bold = not closing
            else:
                italic = not closing
        pos = m.end()
    if pos < len(text):
        runs.append((text[pos:], bold, italic))
    return [r for r in runs if r[0]]


def _pdf_markup(text: str) -> str:
    """Convert inline <b>/<i> runs into ReportLab's own markup, XML-escaping the
    plain-text portions (ReportLab's Paragraph parser natively understands <b>/<i>)."""
    parts = []
    for seg, bold, italic in _split_inline_runs(text):
        seg = xml_escape(seg)
        if bold:
            seg = f"<b>{seg}</b>"
        if italic:
            seg = f"<i>{seg}</i>"
        parts.append(seg)
    return "".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# PDF  (ReportLab Platypus)
# ─────────────────────────────────────────────────────────────────────────────

def _rl_image(data_url: str, max_w: float, max_h: float) -> RLImage | None:
    """Build a centred ReportLab Image flowable, scaled to fit max_w × max_h (pts)."""
    raw = _decode_image(data_url)
    if not raw:
        return None
    try:
        buf = io.BytesIO(raw)
        pil = PILImage.open(buf)
        w, h = pil.size
        if not w or not h:
            return None
        scale = min(max_w / w, max_h / h, 1.0)
        buf.seek(0)
        img = RLImage(buf, width=w * scale, height=h * scale)
        img.hAlign = "CENTER"
        return img
    except Exception:
        return None


def _page_num_cb(canvas, doc) -> None:
    """Footer callback: draws '— N —' centred at the bottom of every page."""
    canvas.saveState()
    canvas.setFont("Times-Roman", 10)
    canvas.setFillColor(colors.HexColor("#888888"))
    canvas.drawCentredString(PAGE_W / 2, 0.42 * inch, f"\u2014 {doc.page} \u2014")
    canvas.restoreState()


def _rich_text_flowables(text: str, S: dict) -> list:
    """Render AI-generated text (headings/bold/italic/lists/code) into ReportLab
    flowables, in place of dumping the raw HTML-tagged string into one Paragraph."""
    out: list = []
    for block in _parse_content_blocks(text):
        if block["type"] == "heading":
            out.append(Paragraph(_pdf_markup(block["text"]), S["subhead"]))
        elif block["type"] == "code":
            out.append(Preformatted(block["text"], S["code"]))
        elif block["type"] == "list":
            for idx, item in enumerate(block["items"]):
                prefix = f"{idx + 1}. " if block["ordered"] else "\u2022 "
                out.append(Paragraph(prefix + _pdf_markup(item), S["list_item"]))
        else:
            out.append(Paragraph(_pdf_markup(block["text"]), S["body"]))
    return out


def generate_pdf(ebook_json: dict, book_title: str, language: str | None = "en") -> bytes:
    """Return raw PDF bytes for the given ebook_json."""

    font_regular, font_bold, font_italic = _fonts_for_language(language)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN + 0.3 * inch,  # extra room for page-number footer
        title=book_title,
        author=ebook_json.get("author", ""),
    )

    # ── Style factory ─────────────────────────────────────────────────────────
    def sty(
        name: str, *,
        bold: bool = False, italic: bool = False, size: int = 12,
        align: int = TA_LEFT, color: str = "#000000",
        before: int = 0, after: int = 8, left_indent: int = 0,
        leading: float | None = None,
    ) -> ParagraphStyle:
        if bold:
            fn = font_bold
        elif italic:
            fn = font_italic
        else:
            fn = font_regular
        return ParagraphStyle(
            name,
            fontName=fn,
            fontSize=size,
            leading=leading if leading is not None else size * 1.5,
            alignment=align,
            spaceBefore=before,
            spaceAfter=after,
            textColor=colors.HexColor(color),
            leftIndent=left_indent,
        )

    S = {
        "cov_title":  sty("ct",  bold=True,   size=30, align=TA_CENTER),
        "cov_sub":    sty("cs",  italic=True, size=14, align=TA_CENTER, color="#444444", after=6),
        "cov_author": sty("ca",                size=14, align=TA_CENTER, color="#222222", after=16),
        "cov_desc":   sty("cd",  italic=True, size=11, align=TA_CENTER, color="#555555"),
        "sec_title":  sty("st",  bold=True,   size=20, align=TA_CENTER, after=18),
        "body":       sty("bd",                size=12, align=TA_JUSTIFY),
        "subhead":    sty("sh",  bold=True,   size=15, before=10, after=6),
        "list_item":  sty("li",                size=12, before=1, after=1, left_indent=16),
        "code":       ParagraphStyle(
                          "code", fontName="Courier", fontSize=9, leading=12,
                          textColor=colors.HexColor("#222222"),
                          backColor=colors.HexColor("#f5f5f5"),
                          borderPadding=8, spaceBefore=6, spaceAfter=10),
        "eyebrow":    sty("ey",                size=9,  color="#777777", after=8),
        "ch_title":   sty("cht", bold=True,   size=22, after=8),
        "toc_entry":  sty("te",                size=12, leading=20),
        "toc_pg":     sty("tp",                size=12, align=TA_RIGHT, color="#555555"),
        "dots":       ParagraphStyle(
                          "dots", fontName="Times-Roman", fontSize=10, leading=20,
                          alignment=TA_CENTER, textColor=colors.HexColor("#cccccc")),
        "kp_label":   sty("kpl", bold=True,   size=9,  color="#333333", after=6),
        "kp_item":    sty("kpi",               size=11, color="#111111", left_indent=14),
        "aq_grp":     sty("aqg", bold=True,   size=14, before=16, after=8),
        "aq_q":       sty("aqq", bold=True,   size=12, after=4),
        "aq_opt":     sty("aqo",               size=11, left_indent=20, after=2),
        "aq_ans":     sty("aqa", italic=True, size=10, color="#555555", left_indent=14, after=8),
        "ty_title":   sty("tyt", bold=True,   size=20, align=TA_CENTER, color="#222222", after=16),
        "ty_text":    sty("tyt2", italic=True, size=12, align=TA_CENTER, color="#444444"),
    }

    # ── Unpack ebook data ─────────────────────────────────────────────────────
    ej = ebook_json
    tp        = ej.get("title_page", {})
    toc       = ej.get("table_of_contents", [])
    chapters  = ej.get("chapters", [])
    ch_imgs   = ej.get("images", {}).get("chapter_images", {})
    cov_img   = ej.get("images", {}).get("cover_image")
    author    = tp.get("author") or ej.get("author", "")
    subtitle  = tp.get("subtitle", "")
    descr     = tp.get("description", "")
    summary   = ej.get("book_summary", "")
    thanks    = ej.get("thank_you_message", "")
    asmnt     = ej.get("final_assessment")

    story: list = []

    # ── Cover ─────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.8 * inch))
    story.append(Paragraph(_pdf_text(book_title), S["cov_title"]))
    if subtitle:
        story.append(Paragraph(_pdf_text(subtitle), S["cov_sub"]))
    story.append(HRFlowable(
        width=60, thickness=2, color=colors.HexColor("#333333"),
        spaceBefore=10, spaceAfter=12,
    ))
    if author:
        story.append(Paragraph(f"by {_pdf_text(author)}", S["cov_author"]))
    if cov_img:
        img = _rl_image(cov_img, CONTENT_W * 0.75, 2.8 * inch)
        if img:
            story += [Spacer(1, 0.1 * inch), img]
    if descr:
        story += [Spacer(1, 0.25 * inch), Paragraph(_pdf_text(descr), S["cov_desc"])]
    story.append(PageBreak())

    # ── About This Book ───────────────────────────────────────────────────────
    if summary:
        story.append(Paragraph("About This Book", S["sec_title"]))
        story.extend(_rich_text_flowables(summary, S))
        story.append(PageBreak())

    # ── Table of Contents ─────────────────────────────────────────────────────
    if toc:
        story.append(Paragraph("Table of Contents", S["sec_title"]))
        story.append(Spacer(1, 0.2 * inch))

        pg_count  = ej.get("page_count", 15)
        front     = 1 + (1 if summary else 0) + 1
        first_pg  = front + 1
        avg       = max(1, round((pg_count - front) / max(len(chapters), 1)))

        toc_rows = []
        for i, item in enumerate(toc):
            num  = str(item.get("chapter_number", i + 1))
            ttl  = item.get("title", "")
            # Use content-aware page_number stored at generation time; fall back to formula
            pg   = str(item.get("page_number") or (first_pg + i * avg))
            toc_rows.append([
                Paragraph(f"<b>{_pdf_text(num)}.</b>\u2002{_pdf_text(ttl)}", S["toc_entry"]),
                Paragraph("." * 50, S["dots"]),
                Paragraph(_pdf_text(pg), S["toc_pg"]),
            ])

        tbl = Table(
            toc_rows,
            colWidths=[CONTENT_W * 0.56, CONTENT_W * 0.30, CONTENT_W * 0.14],
        )
        tbl.setStyle(TableStyle([
            ("VALIGN",         (0, 0), (-1, -1), "BOTTOM"),
            ("TOPPADDING",     (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)
        story.append(PageBreak())

    # ── Chapters ──────────────────────────────────────────────────────────────
    for i, ch in enumerate(chapters):
        ch_num = ch.get("chapter_number", i + 1)
        ch_ttl = ch.get("title", "")
        body   = ch.get("content") or ch.get("description", "")
        kps    = ch.get("key_points") or []
        imgs   = ch_imgs.get(str(i), [])

        ch_story: list = [
            Paragraph(f"CHAPTER {ch_num}", S["eyebrow"]),
            Paragraph(_pdf_text(ch_ttl), S["ch_title"]),
            Spacer(1, 0.15 * inch),
        ]

        # Image 1  (before body text)
        if imgs and imgs[0]:
            img = _rl_image(imgs[0], CONTENT_W * 0.8, 2.6 * inch)
            if img:
                ch_story += [Spacer(1, 0.1 * inch), img, Spacer(1, 0.2 * inch)]

        # Body paragraphs
        ch_story.extend(_rich_text_flowables(body, S))

        # Image 2  (after body text)
        if len(imgs) > 1 and imgs[1]:
            img = _rl_image(imgs[1], CONTENT_W * 0.8, 2.6 * inch)
            if img:
                ch_story += [Spacer(1, 0.2 * inch), img, Spacer(1, 0.2 * inch)]

        # Key-points box
        if kps:
            kp_inner: list = [Paragraph("KEY POINTS", S["kp_label"])]
            for kp in kps:
                # U+2022 (bullet), not U+25B8 (triangle) \u2014 confirmed present in every
                # font used here (Times-Roman and all registered Noto Sans <Script>
                # fonts); U+25B8 isn't, and silently draws as a missing-glyph box.
                kp_inner.append(Paragraph(f"\u2022\u2002{_pdf_markup(kp)}", S["kp_item"]))
            kp_tbl = Table([[kp_inner]], colWidths=[CONTENT_W])
            kp_tbl.setStyle(TableStyle([
                ("LEFTPADDING",   (0, 0), (-1, -1), 14),
                ("RIGHTPADDING",  (0, 0), (-1, -1), 14),
                ("TOPPADDING",    (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor("#f6f6f6")),
                ("LINEBEFORE",    (0, 0), (0, -1),  4, colors.HexColor("#333333")),
            ]))
            ch_story += [Spacer(1, 0.15 * inch), kp_tbl]

        # The preceding section (cover, or "About This Book"/ToC if present) already
        # ends on its own PageBreak(), so only chapters after the first need one —
        # otherwise two consecutive PageBreak()s produce a genuinely blank page.
        if i > 0:
            story.append(PageBreak())
        story.extend(ch_story)

    # ── Assessment ────────────────────────────────────────────────────────────
    if asmnt:
        story.append(PageBreak())
        story.append(Paragraph("Assessment Questions", S["sec_title"]))
        story.append(Spacer(1, 0.2 * inch))

        def _qgroup(qs: list, label: str, qtype: str) -> None:
            if not qs:
                return
            story.append(Paragraph(label, S["aq_grp"]))
            for j, q in enumerate(qs):
                ch_ref = q.get("chapter_number", "")
                blk: list = [
                    Paragraph(
                        f"{j + 1}. {_pdf_markup(q.get('question', ''))} "
                        f"<font color='#aaaaaa' size='9'>(Ch.\u202f{_pdf_text(ch_ref)})</font>",
                        S["aq_q"],
                    )
                ]
                if qtype == "mcq":
                    for k, opt in enumerate(q.get("options") or []):
                        blk.append(Paragraph(f"{chr(65 + k)})\u2002{_pdf_markup(opt)}", S["aq_opt"]))
                if q.get("answer"):
                    blk.append(Paragraph(f"Answer:\u2002{_pdf_markup(str(q['answer']))}", S["aq_ans"]))
                story.append(KeepTogether(blk))
                story.append(Spacer(1, 0.08 * inch))

        _qgroup(asmnt.get("mcq_questions"),          "Multiple Choice Questions", "mcq")
        _qgroup(asmnt.get("fill_in_blank_questions"), "Fill in the Blanks",       "fib")
        _qgroup(asmnt.get("short_answer_questions"),  "Short Answer Questions",   "sa")
        _qgroup(asmnt.get("long_answer_questions"),   "Long Answer Questions",    "la")

    # ── Thank You ─────────────────────────────────────────────────────────────
    if thanks:
        story.append(PageBreak())
        story += [
            Spacer(1, 1.5 * inch),
            Paragraph("Thank You", S["ty_title"]),
            Spacer(1, 0.2 * inch),
            Paragraph(_pdf_markup(thanks), S["ty_text"]),
        ]

    doc.build(story, onFirstPage=_page_num_cb, onLaterPages=_page_num_cb)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX  (python-docx)
# ─────────────────────────────────────────────────────────────────────────────

def _split_paras(text: str) -> list[str]:
    """Split raw text into non-empty paragraphs (separated by blank lines)."""
    return [p.strip().replace("\n", " ") for p in text.split("\n\n") if p.strip()]


def _page_break(doc: Document) -> None:
    """Insert a hard page break into the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    br  = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _set_run_size(run, size_pt: int) -> None:
    """Set a run's font size for BOTH Western and complex scripts.

    run.font.size only writes <w:sz> (the Western-script size). Word renders
    non-Latin scripts (Tamil, Devanagari, Telugu, Kannada, Malayalam, ...) using
    the SEPARATE <w:szCs> (complex-script size) instead, and python-docx has no
    high-level API for it. Left unset, szCs falls back to the Normal style's
    default — so a 22pt chapter title in Tamil would still *display* at whatever
    the body-text default is, even though the file correctly stores 22pt for w:sz.
    Setting both keeps headings visually distinct from body text in every language.
    """
    run.font.size = Pt(size_pt)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), str(size_pt * 2))  # half-points, same unit as w:sz
    run._r.get_or_add_rPr().append(sz_cs)


def _sp(
    doc: Document, text: str, *,
    bold: bool = False, italic: bool = False, size: int = 12,
    align: int = WD_ALIGN_PARAGRAPH.LEFT,
    before: int = 0, after: int = 8,
    color: tuple[int, int, int] | None = None,
    font_name: str = FONT_DOCX,
) -> None:
    """Add a styled paragraph to doc (Times New Roman by default)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(size * 1.5)
    run = p.add_run(text)
    run.font.name   = font_name
    _set_run_size(run, size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _rich_paragraph(
    doc: Document, text: str, *,
    size: int = 12, align: int = WD_ALIGN_PARAGRAPH.LEFT,
    before: int = 0, after: int = 8, base_bold: bool = False, base_italic: bool = False,
    font_name: str = FONT_DOCX, color: tuple[int, int, int] | None = None,
) -> None:
    """Add a paragraph split into multiple runs so inline <b>/<i> tags render as
    real bold/italic instead of literal text (any other stray tag is dropped)."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    p.paragraph_format.line_spacing = Pt(size * 1.5)
    for seg, bold, italic in _split_inline_runs(text):
        run = p.add_run(seg)
        run.font.name   = font_name
        _set_run_size(run, size)
        run.font.bold   = base_bold or bold
        run.font.italic = base_italic or italic
        if color:
            run.font.color.rgb = RGBColor(*color)


def _rich_text_docx(doc: Document, text: str, *, size: int = 12, align: int = WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Render AI-generated text (headings/bold/italic/lists/code) as real Word
    elements, in place of dumping the raw HTML-tagged string into one paragraph."""
    for block in _parse_content_blocks(text):
        if block["type"] == "heading":
            _rich_paragraph(doc, block["text"], size=15, base_bold=True, before=10, after=6)
        elif block["type"] == "code":
            lines = block["text"].split("\n")
            for i, line in enumerate(lines):
                _sp(doc, line or " ", size=10, font_name="Courier New",
                    color=(51, 51, 51), before=4 if i == 0 else 0, after=4 if i == len(lines) - 1 else 0)
        elif block["type"] == "list":
            for idx, item in enumerate(block["items"]):
                prefix = f"{idx + 1}. " if block["ordered"] else "•  "
                _rich_paragraph(doc, prefix + item, size=size, before=1, after=1)
        else:
            _rich_paragraph(doc, block["text"], size=size, align=align, after=8)


def _docx_image(doc: Document, data_url: str, max_w_in: float, max_h_in: float) -> None:
    """Insert a centred image from a base64 data URL (scales to fit constraints)."""
    raw = _decode_image(data_url)
    if not raw:
        return
    try:
        buf   = io.BytesIO(raw)
        pil   = PILImage.open(buf)
        w_px, h_px = pil.size
        if not w_px or not h_px:
            return
        dpi   = 96.0
        scale = min(max_w_in / (w_px / dpi), max_h_in / (h_px / dpi), 1.0)
        buf.seek(0)
        doc.add_picture(buf, width=Inches((w_px / dpi) * scale))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass


def _no_border_cell(cell) -> None:
    """Remove all borders from a table cell."""
    tcPr     = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("left", "right", "top", "bottom", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def _kp_box(doc: Document, key_points: list[str]) -> None:
    """Render a key-points box: gray background + 4pt left border."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]

    # Gray shading
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F6F6F6")
    tcPr.append(shd)

    # Left border only
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tcBorders.append(tag)
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), "333333")
    tcBorders.append(left)
    tcPr.append(tcBorders)

    # "KEY POINTS" label
    lp  = cell.paragraphs[0]
    lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr  = lp.add_run("KEY POINTS")
    lr.font.name  = FONT_DOCX
    lr.font.bold  = True
    _set_run_size(lr, 9)
    lr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Bullet items
    for kp in key_points:
        ip = cell.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for seg, bold, italic in _split_inline_runs(f"\u25b8\u2002{kp}"):
            run = ip.add_run(seg)
            run.font.name   = FONT_DOCX
            _set_run_size(run, 11)
            run.font.bold   = bold
            run.font.italic = italic


def _docx_footer_page_num(doc: Document) -> None:
    """Add centred '— N —' page numbers to all section footers."""
    for section in doc.sections:
        footer = section.footer
        p      = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _run(txt: str) -> None:
            r = p.add_run(txt)
            r.font.name  = FONT_DOCX
            _set_run_size(r, 10)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        _run("\u2014\u2009")  # em-dash + thin space

        # PAGE field code
        r2 = p.add_run()
        r2.font.name = FONT_DOCX
        _set_run_size(r2, 10)
        r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        for tag, ftype in [("begin", None), (None, " PAGE "), ("end", None)]:
            if tag is not None:
                fld = OxmlElement("w:fldChar")
                fld.set(qn("w:fldCharType"), tag)
                r2._r.append(fld)
            else:
                instr = OxmlElement("w:instrText")
                instr.set(qn("xml:space"), "preserve")
                instr.text = ftype
                r2._r.append(instr)

        _run("\u2009\u2014")  # thin space + em-dash


def generate_docx(ebook_json: dict, book_title: str) -> bytes:
    """Return raw DOCX bytes for the given ebook_json."""
    doc = Document()

    # A4 page, 1-inch margins on all sides
    for section in doc.sections:
        section.page_width    = Inches(A4_W_IN)
        section.page_height   = Inches(11.69)
        section.top_margin    = Inches(DOC_MARGIN)
        section.bottom_margin = Inches(DOC_MARGIN)
        section.left_margin   = Inches(DOC_MARGIN)
        section.right_margin  = Inches(DOC_MARGIN)

    # Effective image widths (within margins)
    MAX_IMG_W = A4_W_IN - 2 * DOC_MARGIN   # 6.27 in
    MAX_IMG_H = 3.0                          # in

    # ── Unpack ebook data ─────────────────────────────────────────────────────
    ej       = ebook_json
    tp       = ej.get("title_page", {})
    toc      = ej.get("table_of_contents", [])
    chapters = ej.get("chapters", [])
    ch_imgs  = ej.get("images", {}).get("chapter_images", {})
    cov_img  = ej.get("images", {}).get("cover_image")
    author   = tp.get("author") or ej.get("author", "")
    subtitle = tp.get("subtitle", "")
    descr    = tp.get("description", "")
    summary  = ej.get("book_summary", "")
    thanks   = ej.get("thank_you_message", "")
    asmnt    = ej.get("final_assessment")

    # ── Cover ─────────────────────────────────────────────────────────────────
    _sp(doc, book_title,
        bold=True, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, before=48, after=8)
    if subtitle:
        _sp(doc, subtitle,
            italic=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
    if author:
        _sp(doc, f"by {author}", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=20)
    if cov_img:
        _docx_image(doc, cov_img, MAX_IMG_W * 0.75, 2.8)
    if descr:
        _sp(doc, descr, italic=True, size=11,
            align=WD_ALIGN_PARAGRAPH.CENTER, before=12, color=(85, 85, 85))
    _page_break(doc)

    # ── About This Book ───────────────────────────────────────────────────────
    if summary:
        _sp(doc, "About This Book",
            bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
        _rich_text_docx(doc, summary, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _page_break(doc)

    # ── Table of Contents ─────────────────────────────────────────────────────
    if toc:
        _sp(doc, "Table of Contents",
            bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

        pg_count = ej.get("page_count", 15)
        front    = 1 + (1 if summary else 0) + 1
        first_pg = front + 1
        avg      = max(1, round((pg_count - front) / max(len(chapters), 1)))

        tbl = doc.add_table(rows=0, cols=3)
        for i, item in enumerate(toc):
            row  = tbl.add_row().cells
            for cell in row:
                _no_border_cell(cell)

            num = str(item.get("chapter_number", i + 1))
            ttl = item.get("title", "")
            pg  = str(first_pg + i * avg)

            p0 = row[0].paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r0 = p0.add_run(f"{num}. {ttl}")
            r0.font.name = FONT_DOCX
            _set_run_size(r0, 12)

            p1 = row[1].paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run("." * 25)
            r1.font.name = FONT_DOCX
            _set_run_size(r1, 10)
            r1.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

            p2 = row[2].paragraphs[0]
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = p2.add_run(pg)
            r2.font.name = FONT_DOCX
            _set_run_size(r2, 12)
            r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        _page_break(doc)

    # ── Chapters ──────────────────────────────────────────────────────────────
    for i, ch in enumerate(chapters):
        if i > 0:
            _page_break(doc)

        ch_num = ch.get("chapter_number", i + 1)
        ch_ttl = ch.get("title", "")
        body   = ch.get("content") or ch.get("description", "")
        kps    = ch.get("key_points") or []
        imgs   = ch_imgs.get(str(i), [])

        _sp(doc, f"CHAPTER {ch_num}", size=9, color=(119, 119, 119), after=4)
        _sp(doc, ch_ttl, bold=True, size=22, after=16)

        if imgs and imgs[0]:
            _docx_image(doc, imgs[0], MAX_IMG_W * 0.8, MAX_IMG_H - 0.4)

        _rich_text_docx(doc, body, size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

        if len(imgs) > 1 and imgs[1]:
            _docx_image(doc, imgs[1], MAX_IMG_W * 0.8, MAX_IMG_H - 0.4)

        if kps:
            _kp_box(doc, kps)

    # ── Assessment ────────────────────────────────────────────────────────────
    if asmnt:
        _page_break(doc)
        _sp(doc, "Assessment Questions",
            bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)

        def _qgroup_docx(qs: list, label: str, qtype: str) -> None:
            if not qs:
                return
            _sp(doc, label, bold=True, size=14, before=12, after=6)
            for j, q in enumerate(qs):
                _rich_paragraph(doc, f"{j + 1}. {q.get('question', '')}", size=12, base_bold=True, after=4)
                if qtype == "mcq":
                    for k, opt in enumerate(q.get("options") or []):
                        _rich_paragraph(doc, f"   {chr(65 + k)}) {opt}", size=11, after=2)
                if q.get("answer"):
                    _rich_paragraph(doc, f"Answer: {q['answer']}", size=10, base_italic=True,
                                     color=(85, 85, 85), after=8)

        _qgroup_docx(asmnt.get("mcq_questions"),          "Multiple Choice Questions", "mcq")
        _qgroup_docx(asmnt.get("fill_in_blank_questions"), "Fill in the Blanks",       "fib")
        _qgroup_docx(asmnt.get("short_answer_questions"),  "Short Answer Questions",   "sa")
        _qgroup_docx(asmnt.get("long_answer_questions"),   "Long Answer Questions",    "la")

    # ── Thank You ─────────────────────────────────────────────────────────────
    if thanks:
        _page_break(doc)
        _sp(doc, "Thank You",
            bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, before=72, after=16)
        _rich_paragraph(doc, thanks, size=12, base_italic=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER, color=(68, 68, 68))

    _docx_footer_page_num(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
